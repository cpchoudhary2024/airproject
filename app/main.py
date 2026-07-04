from __future__ import annotations

import json
import logging
import os
import re
import shutil
import time
import uuid
from pathlib import Path

from fastapi import FastAPI, File, Form, HTTPException, Request, UploadFile
from fastapi.responses import FileResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles

from fastapi.responses import JSONResponse, StreamingResponse
from .analysis import analyze_dataset, build_comparison_pdf, build_public_report_pdf

logger = logging.getLogger("pair_analyzer")

APP_ROOT = Path(__file__).resolve().parent
# DATA_DIR env var lets any deployment (HF Spaces, Docker, Vercel) set the writable path.
# Falls back to a local data directory for development.
DATA_ROOT = Path(os.environ.get("DATA_DIR", str(APP_ROOT / "data")))
DATA_ROOT.mkdir(parents=True, exist_ok=True)

app = FastAPI(title="PurpleAir Local Analyzer")

app.mount("/static", StaticFiles(directory=str(APP_ROOT / "static")), name="static")


# Content-Security-Policy: lock the page down to same-origin resources plus the
# one third-party we actually load (Plotly from its CDN). Inline styles/scripts
# used by the template are permitted; no external connect/frame targets.
_CSP = (
    "default-src 'self'; "
    "script-src 'self' 'unsafe-inline' https://cdn.plot.ly; "
    "style-src 'self' 'unsafe-inline'; "
    "img-src 'self' data: blob:; "
    "font-src 'self' data:; "
    "connect-src 'self'; "
    "object-src 'none'; "
    "base-uri 'self'; "
    "frame-ancestors 'none'"
)


@app.middleware("http")
async def _security_headers(request: Request, call_next):
    response = await call_next(request)
    response.headers.setdefault("Content-Security-Policy", _CSP)
    response.headers.setdefault("X-Content-Type-Options", "nosniff")
    response.headers.setdefault("X-Frame-Options", "DENY")
    response.headers.setdefault("Referrer-Policy", "no-referrer")
    response.headers.setdefault(
        "Permissions-Policy", "geolocation=(), microphone=(), camera=()"
    )
    return response


@app.exception_handler(Exception)
async def _unhandled_exception_handler(request: Request, exc: Exception) -> JSONResponse:
    """Catch-all so an unexpected error never leaks a stack trace to the client."""
    err_id = uuid.uuid4().hex[:8]
    logger.exception("[%s] Unhandled error on %s %s: %s", err_id, request.method, request.url.path, exc)
    return JSONResponse(
        status_code=500,
        content={"detail": f"Internal server error (error id: {err_id})"},
    )


_UUID_DIR_RE = re.compile(
    r"^[0-9a-f]{8}-[0-9a-f]{4}-[1-5][0-9a-f]{3}-[89ab][0-9a-f]{3}-[0-9a-f]{12}$",
    re.IGNORECASE,
)


def _is_uuid_dir(dir_path: Path) -> bool:
    return bool(_UUID_DIR_RE.match(dir_path.name))


# Only these characters are ever accepted in a download `kind` — no dots, no
# path separators — which by construction makes traversal via `kind` impossible.
_SAFE_KIND_RE = re.compile(r"^[A-Za-z0-9_-]{1,64}$")

# Maximum accepted upload size (per file). Overridable per-deployment.
_DEFAULT_MAX_UPLOAD_MB = 50


def _max_upload_bytes() -> int:
    mb = _parse_int_env("PAIR_ANALYZER_MAX_UPLOAD_MB")
    if mb is None or mb <= 0:
        mb = _DEFAULT_MAX_UPLOAD_MB
    return mb * 1024 * 1024


def _resolve_job_dir(file_id: str) -> Path:
    """Validate a job id and return its resolved directory, or raise 404.

    Guards against path traversal: the id must be a canonical UUID (matching the
    same pattern used to name job dirs) and the resolved directory must be a
    direct child of DATA_ROOT. Anything else is reported as an unknown id so we
    never disclose whether an out-of-tree path exists.
    """
    if not isinstance(file_id, str) or not _UUID_DIR_RE.match(file_id):
        raise HTTPException(status_code=404, detail="Unknown analysis id")
    data_root = DATA_ROOT.resolve()
    resolved = (data_root / file_id).resolve()
    if resolved.parent != data_root or not resolved.is_dir():
        raise HTTPException(status_code=404, detail="Unknown analysis id")
    return resolved


def _safe_job_file(job_dir: Path, name: str) -> Path | None:
    """Resolve `name` inside `job_dir`, refusing anything that escapes the dir.

    Returns the resolved path only if it stays within `job_dir` and is a regular
    file; otherwise None. Used as defense-in-depth even for server-generated
    filenames so a poisoned summary.json can never point outside the job dir.
    """
    if not name or "\x00" in name:
        return None
    base = job_dir.resolve()
    candidate = (base / name).resolve()
    try:
        candidate.relative_to(base)
    except ValueError:
        return None
    return candidate if candidate.is_file() else None


def _log_and_500(exc: Exception, context: str) -> HTTPException:
    """Log the full error server-side and return a generic client-facing 500.

    Never leaks stack traces or internal exception text to the client; instead
    returns a short correlation id the user can quote in a bug report.
    """
    err_id = uuid.uuid4().hex[:8]
    logger.exception("[%s] %s: %s", err_id, context, exc)
    return HTTPException(status_code=500, detail=f"{context} failed (error id: {err_id})")


def _parse_int_env(var_name: str) -> int | None:
    raw = os.getenv(var_name)
    if raw is None or raw.strip() == "":
        return None

    try:
        return int(raw)
    except ValueError:
        return None


def _env_truthy(var_name: str) -> bool:
    raw = os.getenv(var_name)
    if raw is None:
        return False
    return raw.strip().lower() in {"1", "true", "yes", "on"}


def _cleanup_data_root(*, keep_dirs: set[Path] | None = None) -> None:
    """Optionally delete older analysis job folders under app/data.

    By default this keeps a small number of the most recent UUID job folders to
    prevent unbounded growth. You can override by setting either:
    - PAIR_ANALYZER_MAX_JOBS: keep only the N most-recent job folders
    - PAIR_ANALYZER_MAX_AGE_DAYS: delete job folders older than N days

    Note: this deletes entire job folders, which will make old download links invalid.
    """

    keep_dirs = keep_dirs or set()

    if _env_truthy("PAIR_ANALYZER_DISABLE_CLEANUP"):
        return

    # Default behavior: keep a small recent history (safe for iterative use).
    default_max_jobs = 5

    max_jobs = _parse_int_env("PAIR_ANALYZER_MAX_JOBS")
    max_age_days = _parse_int_env("PAIR_ANALYZER_MAX_AGE_DAYS")

    if max_jobs is None and max_age_days is None:
        max_jobs = default_max_jobs

    if max_jobs is not None and max_jobs < 0:
        return
    if max_age_days is not None and max_age_days < 0:
        return

    job_dirs: list[Path] = [p for p in DATA_ROOT.iterdir() if p.is_dir() and _is_uuid_dir(p)]
    if not job_dirs:
        return

    # Newest first
    job_dirs.sort(key=lambda p: p.stat().st_mtime, reverse=True)
    to_delete: set[Path] = set()

    if max_jobs is not None:
        to_delete.update(job_dirs[max_jobs:])

    if max_age_days is not None:
        cutoff = time.time() - (max_age_days * 86400)
        for job_dir in job_dirs:
            try:
                if job_dir.stat().st_mtime < cutoff:
                    to_delete.add(job_dir)
            except FileNotFoundError:
                continue

    # Never delete explicitly kept dirs.
    to_delete.difference_update(keep_dirs)

    # Delete oldest first (more predictable for logs)
    for job_dir in sorted(to_delete, key=lambda p: p.stat().st_mtime):
        shutil.rmtree(job_dir, ignore_errors=True)


@app.get("/", response_class=HTMLResponse)
def index() -> HTMLResponse:
    index_path = APP_ROOT / "templates" / "index.html"
    return FileResponse(path=str(index_path), media_type="text/html")


async def _run_analysis_job(file: UploadFile, *, generate_outputs: bool = True, metadata: dict | None = None) -> tuple[str, Path, dict, dict]:
    if not file.filename:
        raise HTTPException(status_code=400, detail="Missing file name")
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in {".csv", ".xlsx"}:
        raise HTTPException(status_code=400, detail="Only .csv or .xlsx supported")

    file_id = str(uuid.uuid4())
    job_dir = DATA_ROOT / file_id
    job_dir.mkdir(parents=True, exist_ok=True)

    # Stream the upload to disk in chunks with a hard size cap so a large (or
    # maliciously huge) file can never exhaust memory or fill the disk.
    max_bytes = _max_upload_bytes()
    raw_path = job_dir / f"raw{ext}"
    total = 0
    try:
        with raw_path.open("wb") as f:
            while True:
                chunk = await file.read(1024 * 1024)
                if not chunk:
                    break
                total += len(chunk)
                if total > max_bytes:
                    f.close()
                    shutil.rmtree(job_dir, ignore_errors=True)
                    raise HTTPException(
                        status_code=413,
                        detail=f"File too large (limit {max_bytes // (1024 * 1024)} MB)",
                    )
                f.write(chunk)
    finally:
        await file.close()

    if total == 0:
        shutil.rmtree(job_dir, ignore_errors=True)
        raise HTTPException(status_code=400, detail="Uploaded file is empty")

    try:
        result, outputs = analyze_dataset(raw_path, job_dir, generate_outputs=generate_outputs, metadata=metadata)
    except Exception as exc:
        raise _log_and_500(exc, "Analysis")

    # Store summary payload with outputs for reuse or debugging.
    summary_data = result.copy()
    summary_data["outputs"] = outputs
    with (job_dir / "summary.json").open("w", encoding="utf-8") as f:
        json.dump(summary_data, f, ensure_ascii=True)

    return file_id, job_dir, result, outputs


@app.post("/api/analyze")
async def analyze(
    file: UploadFile = File(...),
    device_id: str = Form(""),
    location: str = Form(""),
    timezone: str = Form("UTC"),
) -> dict:
    metadata = {"device_id": device_id.strip(), "location": location.strip(), "timezone": timezone.strip()}
    file_id, job_dir, result, outputs = await _run_analysis_job(file, metadata=metadata)

    # Store metadata in summary.json for comparison report use
    _cleanup_data_root(keep_dirs={job_dir})

    return {"file_id": file_id, "result": result, "outputs": outputs}


@app.post("/api/analyze-multi")
async def analyze_multi(
    files: list[UploadFile] = File(...),
    labels: str = Form(""),
    control_index: int = Form(0),
) -> dict:
    if not files or len(files) < 2:
        raise HTTPException(status_code=400, detail="Upload at least 2 files")
    if len(files) > 10:
        raise HTTPException(status_code=400, detail="Upload up to 10 files")

    # Parse optional per-file display labels (JSON array of strings, same order as files)
    label_list: list[str] = []
    if labels:
        try:
            parsed = json.loads(labels)
            if isinstance(parsed, list):
                label_list = [str(x) for x in parsed]
        except Exception:
            label_list = []

    analyses: list[dict] = []
    keep_dirs: set[Path] = set()
    _house_counter = 0  # sequential numbering for non-control files

    for idx, file in enumerate(files):
        # Generate full outputs for the Control House so a community report
        # (which needs daily_aqi.csv / hourly_summary.csv + figures) can be built for it.
        _gen = (idx == control_index)
        file_id, job_dir, result, outputs = await _run_analysis_job(file, generate_outputs=_gen)
        keep_dirs.add(job_dir)

        # Resolve a display label: explicit label > "Control House" / "House N" fallback
        if idx == control_index:
            label = (label_list[idx].strip() if idx < len(label_list) and label_list[idx].strip()
                     else "Control House")
        else:
            _house_counter += 1
            label = (label_list[idx].strip() if idx < len(label_list) and label_list[idx].strip()
                     else f"House {_house_counter}")

        # Return a comparison-focused payload (still includes ids for downloads).
        analyses.append(
            {
                "file_id": file_id,
                "filename": file.filename,
                "label": label,
                "is_control": idx == control_index,
                "summary": result.get("summary", {}),
                "timeseries": result.get("charts", {}).get("timeseries", {}),
                "rolling_median": result.get("charts", {}).get("rolling_median", {}),
                "daily_doy": result.get("charts", {}).get("daily_doy", []),
                "outputs": outputs,
            }
        )

    # Difference-in-differences: quantify each house's excess vs the control.
    from .analysis import attach_did_to_analyses
    try:
        attach_did_to_analyses(analyses)
    except Exception as exc:
        logger.warning("DiD computation skipped: %s", exc)

    # Run cleanup once so we never delete jobs from this batch.
    _cleanup_data_root(keep_dirs=keep_dirs)

    return {"analyses": analyses}


@app.get("/api/download/{file_id}/report_word")
def export_to_word(file_id: str) -> FileResponse:
    """Export analysis report as editable Word document (.docx)"""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    import pandas as pd
    
    job_dir = _resolve_job_dir(file_id)
    
    # Load summary.json
    summary_path = job_dir / "summary.json"
    if not summary_path.exists():
        raise HTTPException(status_code=404, detail="Summary data not found")
    
    try:
        with summary_path.open("r", encoding="utf-8") as f:
            summary_data = json.load(f)
    except Exception as e:
        raise _log_and_500(e, "Reading summary")
    
    try:
        # Create Word document
        doc = Document()
        
        # Get summary data
        result = summary_data.get("result", {})
        summary = result.get("summary", summary_data.get("summary", {}))
        detected = result.get("detected", summary_data.get("detected", {}))
        date_range = summary.get("date_range", {})
        
        # Add title
        title = doc.add_heading('Air Quality Analysis Report', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add subtitle with date
        subtitle = doc.add_paragraph('Professional Environmental Data Analysis')
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        subtitle_format = subtitle.runs[0].font
        subtitle_format.size = Pt(12)
        subtitle_format.color.rgb = RGBColor(100, 100, 100)
        
        # Date range
        date_p = doc.add_paragraph()
        date_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        date_p.add_run(f"Analysis Period: {date_range.get('start', 'N/A')} to {date_range.get('end', 'N/A')}")
        
        doc.add_paragraph()  # Spacing
        
        # ===== 1. EXECUTIVE SUMMARY =====
        doc.add_heading('1. Executive Summary', level=1)
        
        summary_section = doc.add_paragraph()
        summary_section.add_run(f"Quality Score: ").bold = True
        summary_section.add_run(f"{summary.get('quality_score', 'N/A')}%")
        
        doc.add_paragraph(f"Period: {date_range.get('start', 'N/A')} to {date_range.get('end', 'N/A')}")
        doc.add_paragraph(f"Quality Score: {summary.get('quality_score', 'N/A')}% | Total Readings: {summary.get('total_readings', 'N/A')}")
        doc.add_paragraph(
            f"Average PM2.5: {summary.get('pm25_average', 'N/A')} µg/m³ | "
            f"AQI: {summary.get('aqi_average', 'N/A')} ({summary.get('aqi_category', 'N/A')})"
        )
        
        if summary.get('quality_narrative'):
            doc.add_paragraph(summary['quality_narrative'], style='List Bullet')
        
        # ===== 2. KEY METRICS TABLE =====
        doc.add_heading('2. Key Metrics Summary', level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Metric'
        hdr_cells[1].text = 'Value'
        
        metrics = [
            ('Current AQI', f"{summary.get('aqi_current', 'N/A')} - {summary.get('aqi_category', 'N/A')}"),
            ('Average AQI', summary.get('aqi_average', 'N/A')),
            ('AQI 10th Percentile', summary.get('aqi_10pct', 'N/A')),
            ('AQI 90th Percentile', summary.get('aqi_90pct', 'N/A')),
            ('Average PM2.5 (µg/m³)', summary.get('pm25_average', 'N/A')),
            ('EPA-Corrected Average', summary.get('pm25_average_epa_corrected', 'N/A')),
            ('PM2.5 10th Percentile', summary.get('pm25_10pct', 'N/A')),
            ('PM2.5 90th Percentile', summary.get('pm25_90pct', 'N/A')),
            ('Quality Score', f"{summary.get('quality_score', 'N/A')}%"),
            ('Validity Score', f"{summary.get('validity_score', 'N/A')}%"),
            ('Coverage Score', f"{summary.get('coverage_score', 'N/A')}%"),
            ('Sensor Health (CV)', f"{summary.get('sensor_health_cv', 'N/A')}%"),
            ('Channel Agreement (R²)', summary.get('channel_agreement_r2', 'N/A')),
            ('Channel Agreement Rate', f"{summary.get('channel_agreement_pct', 'N/A')}%"),
            ('Total Valid Readings', summary.get('total_readings', 'N/A')),
        ]
        
        for metric_name, metric_value in metrics:
            row_cells = table.add_row().cells
            row_cells[0].text = str(metric_name)
            row_cells[1].text = str(metric_value)
        
        doc.add_paragraph()  # Spacing
        
        # ===== 3. VISUALIZATIONS =====
        doc.add_heading('3. Key Visualizations', level=1)
        
        # Add charts if they exist
        charts_to_add = [
            ('fig_radar_quality.png', 'Data Quality Profile Radar', 
             'Multi-dimensional quality assessment evaluating PM2.5 validity, timestamp accuracy, sensor validity, '
             'temporal consistency, and data coverage. Scores ≥80% on all metrics indicate research-grade data reliability.'),
            ('fig_channel_ab.png', 'Channel A vs B Correlation',
             'Validates instrument health by comparing dual-sensor measurements. Shows Pearson correlation coefficient and '
             'mean absolute difference. High agreement (R² > 0.85) indicates excellent data quality.'),
            ('fig_diurnal.png', 'Hourly Diurnal Pattern',
             'PM2.5 aggregated by hour of day showing typical daily cycle. Mean (bold line) with 10th-90th percentile bands. '
             'Peaks align with traffic rush hours and emission source activity.'),
            ('fig_drift.png', 'Sensor Drift Detection',
             'Time-series of channel differences with 7-day rolling median. Flat line near zero indicates excellent agreement. '
             'Persistent slope indicates sensor drift requiring recalibration.'),
            ('fig_rolling_medians.png', 'Rolling Median Trends',
             '24-hour and 7-day moving medians filter noise while preserving signal transitions. Shows underlying air quality trend.'),
            ('fig_radar_pm25_temporal.png', 'PM2.5 Temporal Radar (24-Hour)',
             'Hourly average PM2.5 in polar format with 00:00 at top. Distance from center = concentration. Reveals diurnal patterns.'),
            ('fig_stl_residuals.png', 'STL Decomposition Residuals',
             'Pollution events independent of regular diurnal cycle. Red points highlight anomalies >2σ from baseline. '
             'Identifies non-routine pollution sources.'),
        ]
        
        for filename, chart_title, chart_description in charts_to_add:
            chart_path = job_dir / filename
            if chart_path.exists():
                doc.add_heading(chart_title, level=2)
                doc.add_paragraph(chart_description, style='List Bullet')
                try:
                    doc.add_picture(str(chart_path), width=Inches(5.5))
                except Exception as e:
                    doc.add_paragraph(f"[Chart image: {filename}]", style='List Bullet')
                doc.add_paragraph()  # Spacing
        
        # ===== 4. METHODS & DATA QUALITY =====
        doc.add_heading('4. Methods & Data Quality Framework', level=1)
        
        doc.add_heading('Sensor Validation', level=2)
        doc.add_paragraph(
            'All measurements checked against physically plausible ranges:\n'
            '• PM2.5: 0–1000 µg/m³\n'
            '• Temperature: -40 to +140°F\n'
            '• Humidity: 0–100%\n'
            '• Pressure: 800–1100 hPa',
            style='List Bullet'
        )
        
        doc.add_heading('Data Processing', level=2)
        doc.add_paragraph(
            'Processing pipeline:\n'
            '• Timestamps validated for monotonic progression\n'
            '• Duplicates removed\n'
            '• EPA correction factors applied (0.534×PM + 5.604 − 0.0844×RH when RH available)\n'
            '• Detection limit correction applied (values <1 µg/m³ corrected to ~0.707 µg/m³)',
            style='List Bullet'
        )
        
        doc.add_heading('Quality Scoring', level=2)
        doc.add_paragraph(
            'Quality Score Formula: Q = (0.7 × Validity Score) + (0.3 × Coverage Score)\n'
            'This formula prioritizes data integrity (70% weight) while accounting for temporal gaps and completeness (30% weight).',
            style='List Bullet'
        )
        
        doc.add_heading('Analytical Methods', level=2)
        doc.add_paragraph(
            'Temporal Pattern Analysis: Data aggregated by hour of day to identify diurnal cycles. Mean and percentiles (10th, 90th) '
            'computed to show typical variability.',
            style='List Bullet'
        )
        doc.add_paragraph(
            'Trend Decomposition: STL decomposition isolates trend (monotonic direction), seasonality (24-hour cycle), and residuals '
            '(anomalies). Distinguishes cyclical patterns from baseline shifts.',
            style='List Bullet'
        )
        doc.add_paragraph(
            'Rolling Medians: 24-hour and 7-day moving medians applied to filter short-term noise while preserving signal transitions. '
            'Median (vs mean) resists extreme outliers.',
            style='List Bullet'
        )
        
        # ===== 5. SENSOR PERFORMANCE =====
        doc.add_heading('5. Sensor Performance Analysis', level=1)
        
        if summary.get('channel_agreement_r2'):
            doc.add_paragraph(f"Channel Agreement (R²): {summary.get('channel_agreement_r2', 'N/A')}")
            doc.add_paragraph(f"Mean Absolute Difference: {summary.get('channel_agreement_mad', 'N/A')} µg/m³")
            doc.add_paragraph(f"Agreement Rate: {summary.get('channel_agreement_pct', 'N/A')}%")
            
            r2_val = float(str(summary.get('channel_agreement_r2', '0')).split()[0])
            if r2_val > 0.85:
                status_para = doc.add_paragraph()
                status_para.add_run("✓ Status: Excellent performance - channels are highly consistent.").font.bold = True
            elif r2_val > 0.70:
                status_para = doc.add_paragraph()
                status_para.add_run("⚠ Status: Good performance - minor drift detected.").font.bold = True
            else:
                status_para = doc.add_paragraph()
                status_para.add_run("✗ Status: Poor channel agreement - recalibration recommended.").font.bold = True
        
        doc.add_paragraph(f"Sensor Health (Coefficient of Variation): {summary.get('sensor_health_cv', 'N/A')}%")
        if summary.get('sensor_health_cv'):
            try:
                cv = float(str(summary.get('sensor_health_cv', '0')).split()[0])
                if cv < 10:
                    doc.add_paragraph("✓ Excellent consistency (CV < 10%)", style='List Bullet')
                elif cv < 15:
                    doc.add_paragraph("✓ Acceptable consistency (CV 10-15%)", style='List Bullet')
                else:
                    doc.add_paragraph("⚠ High variability detected (CV > 15%)", style='List Bullet')
            except:
                pass
        
        # ===== 6. DETECTED DATA COLUMNS =====
        if detected:
            doc.add_heading('6. Detected Data Columns', level=1)
            for param, info in detected.items():
                if isinstance(info, dict):
                    p = doc.add_paragraph()
                    p.add_run(f"{param.upper()}: ").bold = True
                    if isinstance(info.get('primary'), dict):
                        p.add_run(
                            f"{info['primary'].get('name', 'Unknown')} "
                            f"(Confidence: {info['primary'].get('confidence', 'N/A')})"
                        )
        
        # ===== 7. REGULATORY COMPARISON =====
        doc.add_heading('7. Comparison to Standards', level=1)
        doc.add_paragraph(
            'EPA 24-Hour Standard: 35 µg/m³ - Unhealthy for sensitive groups when exceeded.',
            style='List Bullet'
        )
        doc.add_paragraph(
            'WHO Guideline: 15 µg/m³ - More stringent target for long-term health protection.',
            style='List Bullet'
        )
        
        avg_pm25 = summary.get('pm25_average', 0)
        if avg_pm25:
            try:
                avg_val = float(str(avg_pm25).split()[0])
                if avg_val > 35:
                    doc.add_paragraph(
                        f"⚠ Alert: Average PM2.5 ({avg_val:.1f} µg/m³) exceeds EPA 24-hour standard.",
                        style='List Bullet'
                    )
                elif avg_val > 15:
                    doc.add_paragraph(
                        f"⚠ Notice: Average PM2.5 ({avg_val:.1f} µg/m³) exceeds WHO guideline.",
                        style='List Bullet'
                    )
            except:
                pass
        
        # ===== 8. KEY FINDINGS =====
        doc.add_heading('8. Key Findings & Anomalies', level=1)
        
        # Try to load anomalies from pollution_events
        try:
            events_path = job_dir / "pollution_events.csv"
            if events_path.exists():
                events_df = pd.read_csv(events_path)
                if len(events_df) > 0:
                    doc.add_paragraph(f"Detected {len(events_df)} pollution events:", style='List Bullet')
                    for idx, row in events_df.head(10).iterrows():
                        event_desc = f"Event {idx+1}: {row.get('timestamp', 'Unknown time')} - {row.get('note', 'Pollution spike detected')}"
                        doc.add_paragraph(event_desc, style='List Bullet')
                    if len(events_df) > 10:
                        doc.add_paragraph(f"... and {len(events_df) - 10} more events (see pollution_events.csv)", style='List Bullet')
                else:
                    doc.add_paragraph("No major pollution events detected.", style='List Bullet')
        except Exception as e:
            doc.add_paragraph("Data quality assessment completed without anomaly detection.", style='List Bullet')
        
        # Add gap analysis if available
        if summary.get('gap_type'):
            doc.add_heading('Data Gaps Analysis', level=2)
            doc.add_paragraph(f"Gap Pattern: {summary.get('gap_type', 'Unknown')}", style='List Bullet')
            doc.add_paragraph(f"Maximum Contiguous Gap: {summary.get('max_contiguous_gap', 'N/A')} hours", style='List Bullet')
            doc.add_paragraph(
                'Note: Contiguous gaps (long stretches) are more damaging to trend decomposition than '
                'stochastic gaps (random blips), as they break temporal sampling assumptions.',
                style='List Bullet'
            )
        
        doc.add_paragraph()  # Spacing
        
        # ===== FOOTER =====
        doc.add_paragraph()
        footer = doc.add_paragraph(
            "Report generated by PurpleAir Local Analyzer - A professional environmental data analysis platform. "
            "Data quality and analysis methodologies adhere to EPA guidance and peer-reviewed environmental science standards."
        )
        footer_format = footer.runs[0].font
        footer_format.size = Pt(9)
        footer_format.color.rgb = RGBColor(150, 150, 150)
        footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Save Word document
        word_path = job_dir / "report_analysis.docx"
        doc.save(word_path)
        
        return FileResponse(path=str(word_path), filename="report_analysis.docx")
        
    except Exception as e:
        raise _log_and_500(e, "Word export")


@app.post("/api/community_report/{file_id}")
async def regenerate_community_report(
    file_id: str,
    device_id: str = Form(""),
    location: str = Form(""),
) -> StreamingResponse:
    """Regenerate the community report with custom device_id / location metadata."""
    import io, tempfile, pandas as pd
    from pathlib import Path as _P

    job_dir = _resolve_job_dir(file_id)
    summary_path = job_dir / "summary.json"
    if not summary_path.exists():
        raise HTTPException(status_code=404, detail="Analysis data not found")
    try:
        with summary_path.open("r", encoding="utf-8") as f:
            stored = json.load(f)
    except Exception:
        raise HTTPException(status_code=500, detail="Could not read analysis data")

    pdf_params  = stored.get("_pdf_params", {})
    base_meta   = pdf_params.get("metadata") or {}
    meta = {
        **base_meta,
        "device_id": device_id.strip() or base_meta.get("device_id", ""),
        "location":  location.strip()  or base_meta.get("location",  ""),
    }

    summary     = stored.get("summary", {})
    # Exceedance counts live under result["tables"]["exceedances"]; fall back to a
    # top-level key for forward compatibility.
    exceedances = stored.get("exceedances") or stored.get("tables", {}).get("exceedances", {})
    anomalies   = stored.get("anomalies", [])
    channel_agreement = pdf_params.get("channel_agreement") or {}

    # Load daily/hourly data — use the actual filenames written by _write_outputs
    daily_path  = job_dir / "daily_aqi.csv"
    hourly_path = job_dir / "hourly_summary.csv"
    if not daily_path.exists() or not hourly_path.exists():
        raise HTTPException(status_code=404, detail="Daily/hourly CSV data not found in job directory")
    daily  = pd.read_csv(daily_path)
    hourly = pd.read_csv(hourly_path)

    # Reconstruct figures using the stored title→filename map saved during analysis
    stored_figs = pdf_params.get("figures") or []
    if stored_figs:
        figures = [
            (title, job_dir / fname)
            for title, fname in stored_figs
            if (job_dir / fname).exists()
        ]
    else:
        # Fallback: known mapping for figures created by build_report_figures
        _known = {
            "Rolling medians":    "fig_rolling_medians.png",
            "Diurnal pattern":    "fig_diurnal.png",
            "PM2.5 Temporal Radar": "fig_radar_pm25_temporal.png",
        }
        figures = [
            (title, job_dir / fname)
            for title, fname in _known.items()
            if (job_dir / fname).exists()
        ]

    tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
    tmp.close()
    try:
        build_public_report_pdf(
            _P(tmp.name), summary, daily, hourly, anomalies,
            channel_agreement, figures, exceedances=exceedances, metadata=meta,
        )
        with open(tmp.name, "rb") as fh:
            content = fh.read()
    finally:
        import os as _os
        try: _os.unlink(tmp.name)
        except Exception: pass

    return StreamingResponse(
        io.BytesIO(content),
        media_type="application/pdf",
        headers={"Content-Disposition": "attachment; filename=community_air_quality_report.pdf"},
    )


@app.post("/api/community_report_compare")
async def community_report_with_comparison(payload: dict) -> StreamingResponse:
    """Community report for the Control House with an appended House Comparison page.

    payload: { file_ids: [...control-first...], labels: [...], device_id, location }
    The first file_id must be the Control House (it has full outputs on disk).
    """
    import io, tempfile, pandas as pd
    from pathlib import Path as _P

    file_ids = payload.get("file_ids") or []
    labels   = payload.get("labels") or []
    if len(file_ids) < 2:
        raise HTTPException(status_code=400, detail="Provide the Control House plus at least one other house")

    control_id = file_ids[0]
    job_dir = _resolve_job_dir(control_id)
    summary_path = job_dir / "summary.json"
    if not summary_path.exists():
        raise HTTPException(status_code=404, detail="Control House analysis not found")
    daily_path  = job_dir / "daily_aqi.csv"
    hourly_path = job_dir / "hourly_summary.csv"
    if not daily_path.exists() or not hourly_path.exists():
        raise HTTPException(status_code=404, detail="Control House output data not found (re-run the comparison)")

    with summary_path.open("r", encoding="utf-8") as f:
        stored = json.load(f)

    pdf_params = stored.get("_pdf_params", {})
    base_meta  = pdf_params.get("metadata") or {}
    meta = {
        **base_meta,
        "device_id": str(payload.get("device_id", "")).strip() or base_meta.get("device_id", ""),
        "location":  str(payload.get("location", "")).strip()  or base_meta.get("location",  ""),
    }
    summary     = stored.get("summary", {})
    # Exceedance counts live under result["tables"]["exceedances"]; fall back to a
    # top-level key for forward compatibility.
    exceedances = stored.get("exceedances") or stored.get("tables", {}).get("exceedances", {})
    anomalies   = stored.get("anomalies", [])
    channel_agreement = pdf_params.get("channel_agreement") or {}
    daily  = pd.read_csv(daily_path)
    hourly = pd.read_csv(hourly_path)

    stored_figs = pdf_params.get("figures") or []
    figures = [(t, job_dir / fn) for t, fn in stored_figs if (job_dir / fn).exists()]

    # Build the comparison list (each house's rolling_median from its summary.json)
    comparison = []
    for i, fid in enumerate(file_ids):
        try:
            sp = _resolve_job_dir(fid) / "summary.json"
        except HTTPException:
            continue
        if not sp.exists():
            continue
        try:
            with sp.open("r", encoding="utf-8") as f:
                d = json.load(f)
        except Exception:
            continue
        comparison.append({
            "label": labels[i] if i < len(labels) and str(labels[i]).strip() else (f"Control House" if i == 0 else f"House {i}"),
            "is_control": i == 0,
            "rolling_median": d.get("charts", {}).get("rolling_median", {}),
        })

    tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
    tmp.close()
    try:
        build_public_report_pdf(
            _P(tmp.name), summary, daily, hourly, anomalies,
            channel_agreement, figures, exceedances=exceedances, metadata=meta,
            comparison=comparison,
        )
        with open(tmp.name, "rb") as fh:
            content = fh.read()
    finally:
        import os as _os
        try: _os.unlink(tmp.name)
        except Exception: pass

    return StreamingResponse(
        io.BytesIO(content),
        media_type="application/pdf",
        headers={"Content-Disposition": "attachment; filename=community_report_with_comparison.pdf"},
    )


@app.get("/api/download/{file_id}/{kind}")
def download(file_id: str, kind: str) -> FileResponse:
    job_dir = _resolve_job_dir(file_id)

    # `kind` is restricted to a safe charset (no dots, no separators) so it can
    # never encode a path. Every candidate is additionally resolved through
    # _safe_job_file, which refuses anything outside the job directory.
    if not _SAFE_KIND_RE.match(kind):
        raise HTTPException(status_code=404, detail="Download not found")

    # First, try to look up the actual filename from summary.json outputs mapping
    summary_path = job_dir / "summary.json"
    if summary_path.exists():
        try:
            with summary_path.open("r", encoding="utf-8") as f:
                summary = json.load(f)
            if isinstance(summary.get("outputs"), dict) and kind in summary["outputs"]:
                file_path = _safe_job_file(job_dir, str(summary["outputs"][kind]))
                if file_path is not None:
                    return FileResponse(path=str(file_path), filename=file_path.name)
        except Exception:
            pass  # Fall through to other methods

    # Try multiple file extensions (pdf, csv, json, docx, xlsx)
    for ext in ['.pdf', '.docx', '.csv', '.json', '.xlsx']:
        file_path = _safe_job_file(job_dir, f"{kind}{ext}")
        if file_path is not None:
            return FileResponse(path=str(file_path), filename=file_path.name)

    # If not found with extensions, try exact name
    file_path = _safe_job_file(job_dir, kind)
    if file_path is not None:
        return FileResponse(path=str(file_path), filename=file_path.name)

    raise HTTPException(status_code=404, detail="Download not found")


@app.post("/api/refine-analysis/{file_id}")
async def refine_analysis(file_id: str, date_from: str, date_to: str) -> dict:
    """Reanalyze data for a specific timeframe."""
    import pandas as pd
    
    job_dir = _resolve_job_dir(file_id)
    
    # Find the raw data file
    raw_files = list(job_dir.glob("raw.*"))
    if not raw_files:
        raise HTTPException(status_code=404, detail="Original data file not found")
    
    raw_path = raw_files[0]
    
    try:
        # Parse date strings to timestamps
        ts_from = pd.Timestamp(date_from)
        ts_to = pd.Timestamp(date_to)
        window = (ts_from, ts_to)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid date format")

    if pd.isna(ts_from) or pd.isna(ts_to):
        raise HTTPException(status_code=400, detail="Invalid date format")

    # Build the subdirectory name from the *parsed* timestamps (never raw user
    # input) so the folder name can't smuggle path separators or "..".
    date_suffix = f"{ts_from.strftime('%Y%m%dT%H')}_{ts_to.strftime('%Y%m%dT%H')}"
    refine_dir = job_dir / f"refined_{date_suffix}"
    refine_dir.mkdir(parents=True, exist_ok=True)

    # Run analysis with the specified time window
    try:
        result, outputs = analyze_dataset(raw_path, job_dir, window=window, output_dir=refine_dir)

        # Store the refined analysis result
        summary_data = result.copy()
        summary_data["outputs"] = outputs
        summary_data["refinement_window"] = {"start": date_from, "end": date_to}
        with (refine_dir / "summary.json").open("w", encoding="utf-8") as f:
            json.dump(summary_data, f, ensure_ascii=True)

        return {"file_id": file_id, "result": result, "outputs": outputs, "refine_dir": str(refine_dir.name)}
    except Exception as e:
        raise _log_and_500(e, "Refinement")


@app.post("/api/generate-report/{file_id}")
async def generate_custom_report(file_id: str, payload: dict) -> FileResponse:
    """Regenerate PDF with custom notes and stored analysis parameters."""
    from .analysis import build_report_pdf

    job_dir = _resolve_job_dir(file_id)

    summary_path = job_dir / "summary.json"
    if not summary_path.exists():
        raise HTTPException(status_code=404, detail="Summary data not found")

    try:
        with summary_path.open("r", encoding="utf-8") as f:
            stored = json.load(f)
    except Exception as e:
        raise _log_and_500(e, "Reading summary")

    custom_notes = payload.get("custom_notes") or []

    # Retrieve stored PDF params
    pdf_params = stored.get("_pdf_params", {})
    tz_label = pdf_params.get("tz_label", "UTC")
    channel_agreement = pdf_params.get("channel_agreement")
    gap_analysis = pdf_params.get("gap_analysis", {})
    metadata = pdf_params.get("metadata") or {}

    summary = stored.get("summary", {})

    # quality_rows is stored under tables.quality
    quality_rows = (stored.get("tables") or {}).get("quality") or []

    # anomalies are at top level
    anomalies = stored.get("anomalies") or []

    # radar_profile is stored under charts.radar_pattern
    radar_profile = (stored.get("charts") or {}).get("radar_pattern")

    # highest_events is stored under tables.highest_events
    highest_events = (stored.get("tables") or {}).get("highest_events") or []

    # Figures must be List[Tuple[str, Path]] — same format build_report_figures returns.
    # Prefer the title→filename map persisted during analysis (always current, includes
    # every figure such as Weekly heatmap and Bland-Altman); fall back to a known list.
    stored_figs = pdf_params.get("figures") or []
    if stored_figs:
        figures = [(title, job_dir / fn) for title, fn in stored_figs if (job_dir / fn).exists()]
    else:
        figure_map = [
            ("Rolling medians",     "fig_rolling_medians.png"),
            ("STL Residuals",       "fig_stl_residuals.png"),
            ("Diurnal pattern",     "fig_diurnal.png"),
            ("PM2.5 Temporal Radar","fig_radar_pm25_temporal.png"),
            ("Weekly heatmap",      "fig_weekly_heatmap.png"),
            ("Channel A vs B",      "fig_channel_ab.png"),
            ("Bland-Altman",        "fig_bland_altman.png"),
            ("Sensor drift",        "fig_drift.png"),
        ]
        figures = [(title, job_dir / fn) for title, fn in figure_map if (job_dir / fn).exists()]

    report_path = job_dir / "report_custom.pdf"
    try:
        build_report_pdf(
            report_path,
            summary,
            quality_rows,
            anomalies,
            figures,
            channel_agreement=channel_agreement,
            gap_analysis=gap_analysis,
            radar_profile=radar_profile,
            metadata=metadata,
            custom_notes=custom_notes,
            tz_label=tz_label,
            highest_events=highest_events,
        )
    except Exception as e:
        raise _log_and_500(e, "PDF generation")

    return FileResponse(path=str(report_path), filename="report_analysis.pdf", media_type="application/pdf")


@app.post("/api/compare-report")
async def generate_compare_report(payload: dict) -> FileResponse:
    """Generate and download a comprehensive comparison PDF report."""
    file_ids: list[str] = payload.get("file_ids", [])
    filenames: list[str] = payload.get("filenames", [])
    labels: list[str] = payload.get("labels", []) or filenames

    if len(file_ids) < 2:
        raise HTTPException(status_code=400, detail="Provide at least 2 file_ids")

    analyses: list[dict] = []
    for i, fid in enumerate(file_ids):
        job_dir = _resolve_job_dir(fid)
        summary_path = job_dir / "summary.json"
        if not summary_path.exists():
            raise HTTPException(status_code=404, detail="Analysis not found")
        try:
            with summary_path.open("r", encoding="utf-8") as f:
                data = json.load(f)
            # summary.json stores result at top level (no nested "result" key).
            # Include rolling_median so the PDF can draw 1h/24h comparison charts.
            analyses.append({
                "summary": data.get("summary", {}),
                "rolling_median": data.get("charts", {}).get("rolling_median", {}),
                "label": labels[i] if i < len(labels) and str(labels[i]).strip() else f"House {i}",
                "is_control": i == 0,  # file_ids arrive control-first from the UI
            })
        except Exception as e:
            raise _log_and_500(e, "Reading analysis")

    if not filenames or len(filenames) < len(file_ids):
        filenames = [a["label"] for a in analyses]

    # Write comparison PDF to a temp dir associated with the first job
    out_dir = _resolve_job_dir(file_ids[0])
    report_path = out_dir / "comparison_report.pdf"
    try:
        build_comparison_pdf(report_path, analyses, filenames)
    except Exception as e:
        raise _log_and_500(e, "PDF generation")

    return FileResponse(path=str(report_path), filename="comparison_report.pdf", media_type="application/pdf")


# ─────────────────────────────────────────────────────────────────────────────
# Opt-in network features (Tier 1). OFF by default: no outbound request is made
# until the user explicitly calls these routes. Only the sensor's coordinates
# and date window are sent to fixed external services — never PM2.5 data.
# ─────────────────────────────────────────────────────────────────────────────

def _load_summary(job_dir: Path) -> dict:
    summary_path = job_dir / "summary.json"
    if not summary_path.exists():
        raise HTTPException(status_code=404, detail="Analysis summary not found")
    try:
        with summary_path.open("r", encoding="utf-8") as f:
            return json.load(f)
    except Exception as e:
        raise _log_and_500(e, "Reading summary")


def _sensor_coords(summary: dict) -> tuple[float, float]:
    s = summary.get("summary", {}) if "summary" in summary else summary
    lat, lon = s.get("latitude"), s.get("longitude")
    if lat is None or lon is None:
        raise HTTPException(
            status_code=422,
            detail="This dataset has no location (latitude/longitude) columns, so location-based features are unavailable.",
        )
    return float(lat), float(lon)


def _load_sensor_hourly_utc(job_dir: Path):
    """Return (hour_key -> pm25) dict from the UTC hourly output, or raise 404."""
    import pandas as pd

    path = _safe_job_file(job_dir, "sensor_hourly_utc.csv")
    if path is None:
        raise HTTPException(status_code=404, detail="Hourly sensor series unavailable for this analysis. Re-run the analysis.")
    df = pd.read_csv(path)
    series: dict[str, float] = {}
    for _, row in df.iterrows():
        key = str(row.get("utc_hour", ""))[:13]
        val = row.get("pm25_corrected")
        if key and pd.notna(val):
            series[key] = float(val)
    if not series:
        raise HTTPException(status_code=404, detail="Hourly sensor series is empty.")
    return series


@app.post("/api/reference-validation/{file_id}")
def reference_validation(file_id: str) -> dict:
    """Compare the sensor against the nearest regulatory PM2.5 monitor (OpenAQ).

    Opt-in: only the sensor's coordinates + date window are sent to OpenAQ.
    Returns collocation statistics (bias, R², RMSE, slope) + paired points."""
    import numpy as np
    from .external import ExternalError, fetch_openaq_reference

    job_dir = _resolve_job_dir(file_id)
    summary = _load_summary(job_dir)
    lat, lon = _sensor_coords(summary)
    sensor = _load_sensor_hourly_utc(job_dir)

    hours = sorted(sensor.keys())
    start_date, end_date = hours[0][:10], hours[-1][:10]

    try:
        ref = fetch_openaq_reference(lat, lon, start_date, end_date)
    except ExternalError as exc:
        raise HTTPException(status_code=502, detail=str(exc))

    ref_series = ref["series"]
    common = [h for h in hours if h in ref_series]
    if len(common) < 8:
        raise HTTPException(
            status_code=422,
            detail=f"Only {len(common)} overlapping hours with the reference monitor — too few for a reliable comparison.",
        )
    s_vals = np.array([sensor[h] for h in common], dtype=float)
    r_vals = np.array([ref_series[h] for h in common], dtype=float)

    diff = s_vals - r_vals
    bias = float(np.mean(diff))
    rmse = float(np.sqrt(np.mean(diff ** 2)))
    mae = float(np.mean(np.abs(diff)))
    if np.std(s_vals) > 0 and np.std(r_vals) > 0:
        r = float(np.corrcoef(s_vals, r_vals)[0, 1])
        r2 = r * r
        slope, intercept = np.polyfit(r_vals, s_vals, 1)
    else:
        r2, slope, intercept = 0.0, float("nan"), float("nan")

    return {
        "monitor": {
            "name": ref.get("location_name"),
            "provider": ref.get("provider"),
            "distance_km": round(ref["distance_m"] / 1000.0, 2) if ref.get("distance_m") is not None else None,
            "latitude": ref.get("latitude"),
            "longitude": ref.get("longitude"),
        },
        "stats": {
            "n_hours": len(common),
            "bias": round(bias, 2),
            "rmse": round(rmse, 2),
            "mae": round(mae, 2),
            "r2": round(r2, 3),
            "slope": round(float(slope), 3) if slope == slope else None,
            "intercept": round(float(intercept), 3) if intercept == intercept else None,
        },
        "scatter": {"sensor": [round(v, 1) for v in s_vals.tolist()],
                    "reference": [round(v, 1) for v in r_vals.tolist()]},
        "timeseries": {
            "hours": common,
            "sensor": [round(sensor[h], 1) for h in common],
            "reference": [round(ref_series[h], 1) for h in common],
        },
        "disclosure": "Only your sensor coordinates and date range were sent to OpenAQ; no PM2.5 data left this server.",
    }


@app.post("/api/pollution-rose/{file_id}")
def pollution_rose(file_id: str) -> dict:
    """Pollution rose: join Open-Meteo wind to the sensor's PM2.5 and bin by
    wind direction to show which direction pollution arrives from.

    Opt-in: only coordinates + date window are sent to Open-Meteo (no key)."""
    import numpy as np
    from .external import ExternalError, fetch_openmeteo_wind

    job_dir = _resolve_job_dir(file_id)
    summary = _load_summary(job_dir)
    lat, lon = _sensor_coords(summary)
    sensor = _load_sensor_hourly_utc(job_dir)

    hours = sorted(sensor.keys())
    start_date, end_date = hours[0][:10], hours[-1][:10]

    try:
        wind = fetch_openmeteo_wind(lat, lon, start_date, end_date)
    except ExternalError as exc:
        raise HTTPException(status_code=502, detail=str(exc))

    # 16-point compass sectors.
    labels = ["N", "NNE", "NE", "ENE", "E", "ESE", "SE", "SSE",
              "S", "SSW", "SW", "WSW", "W", "WNW", "NW", "NNW"]
    n_sec = 16
    sums = [0.0] * n_sec
    counts = [0] * n_sec
    speeds = [0.0] * n_sec

    matched = 0
    for i, t in enumerate(wind["time"]):
        key = str(t)[:13]
        pm = sensor.get(key)
        if pm is None:
            continue
        try:
            direction = float(wind["direction"][i])
            speed = float(wind["speed"][i])
        except (TypeError, ValueError, IndexError):
            continue
        if direction != direction:  # NaN
            continue
        # Sector centered on each compass point (N spans 348.75–11.25°).
        sec = int(((direction % 360) + 11.25) // 22.5) % n_sec
        sums[sec] += pm
        speeds[sec] += speed
        counts[sec] += 1
        matched += 1

    if matched < 8:
        raise HTTPException(
            status_code=422,
            detail=f"Only {matched} hours had both wind and PM2.5 data — too few to build a pollution rose.",
        )

    mean_pm = [round(sums[i] / counts[i], 2) if counts[i] else 0.0 for i in range(n_sec)]
    mean_speed = [round(speeds[i] / counts[i], 1) if counts[i] else 0.0 for i in range(n_sec)]
    total_pm = sum(sums)
    weighted_pct = [round(100.0 * sums[i] / total_pm, 1) if total_pm > 0 else 0.0 for i in range(n_sec)]
    dominant = int(np.argmax(sums)) if total_pm > 0 else None

    return {
        "sectors": labels,
        "angles_deg": [i * 22.5 for i in range(n_sec)],
        "mean_pm25": mean_pm,
        "mean_wind_speed": mean_speed,
        "hours_per_sector": counts,
        "pm_share_pct": weighted_pct,
        "dominant_sector": labels[dominant] if dominant is not None else None,
        "wind_units": wind.get("units", "km/h"),
        "n_hours": matched,
        "disclosure": "Only your sensor coordinates and date range were sent to Open-Meteo; no PM2.5 data left this server.",
    }
